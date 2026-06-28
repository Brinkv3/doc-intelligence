"""Document ingestion.

Parsers for each supported format return uniform Document objects.
Adapted from the RAG pipeline — same parsing logic, same Document dataclass.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".md", ".txt", ".csv"}


@dataclass
class Document:
    """A parsed source document with provenance metadata."""

    doc_id: str
    text: str
    metadata: dict = field(default_factory=dict)


def parse_file(path: str | Path) -> Document:
    """Parse a single file into a Document."""
    path = Path(path)
    if not path.is_file():
        raise FileNotFoundError(f"Not a file: {path}")
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")
    return _parse_file(path)


def parse_directory(corpus_path: str | Path) -> list[Document]:
    """Walk a directory and parse every supported file into Documents."""
    corpus = Path(corpus_path)
    if not corpus.is_dir():
        raise ValueError(f"Corpus path is not a directory: {corpus}")

    documents: list[Document] = []
    for path in sorted(corpus.rglob("*")):
        if not path.is_file():
            continue
        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            logger.info("Skipping unsupported file: %s", path.name)
            continue
        try:
            doc = _parse_file(path)
            if doc.text.strip():
                documents.append(doc)
                logger.info("Parsed %s", path.name)
        except Exception:
            logger.exception("Failed to parse %s — skipping", path.name)

    logger.info("Ingested %d document(s) from %s", len(documents), corpus)
    return documents


def _parse_file(path: Path) -> Document:
    ext = path.suffix.lower()
    parsers = {
        ".docx": _parse_docx,
        ".pdf": _parse_pdf,
        ".md": _parse_text,
        ".txt": _parse_text,
        ".csv": _parse_csv,
    }
    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"No parser for {ext}")
    return parser(path)


def _parse_docx(path: Path) -> Document:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading 1" in style:
            parts.append(f"\n# {text}")
        elif "Heading 2" in style:
            parts.append(f"\n## {text}")
        elif "Heading 3" in style:
            parts.append(f"\n### {text}")
        else:
            parts.append(text)

    for i, table in enumerate(doc.tables):
        table_rows: list[str] = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                table_rows.append(" | ".join(cells))
        if table_rows:
            parts.append(f"\n[Table {i + 1}]\n" + "\n".join(table_rows))

    return Document(
        doc_id=path.name,
        text="\n".join(parts),
        metadata={
            "source_file": path.name,
            "source_path": str(path),
            "doc_type": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        },
    )


def _parse_pdf(path: Path) -> Document:
    import fitz

    pdf = fitz.open(str(path))
    pages = []
    for i, page in enumerate(pdf):
        text = page.get_text().strip()
        if text:
            pages.append(f"[Page {i + 1}]\n{text}")
    pdf.close()

    return Document(
        doc_id=path.name,
        text="\n\n".join(pages),
        metadata={
            "source_file": path.name,
            "source_path": str(path),
            "doc_type": "pdf",
            "page_count": len(pages),
        },
    )


def _parse_text(path: Path) -> Document:
    text = path.read_text(encoding="utf-8")
    return Document(
        doc_id=path.name,
        text=text,
        metadata={
            "source_file": path.name,
            "source_path": str(path),
            "doc_type": path.suffix.lstrip("."),
        },
    )


def _parse_csv(path: Path) -> Document:
    import pandas as pd

    df = pd.read_csv(path)
    text_blocks: list[str] = []
    for _, row in df.iterrows():
        pairs = [f"{col}: {row[col]}" for col in df.columns if pd.notna(row[col])]
        if pairs:
            text_blocks.append("\n".join(pairs))

    return Document(
        doc_id=path.name,
        text="\n\n".join(text_blocks),
        metadata={
            "source_file": path.name,
            "source_path": str(path),
            "doc_type": "csv",
            "row_count": len(text_blocks),
        },
    )
